import docx

file=docx.Document("demo.docx")
print("段落数:"+str(len(file.paragraphs))) #输出段落数
file_word = docx.Document()

#输出每一段的内容
for para in file.paragraphs:
    print(para.text)
